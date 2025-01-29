from flask import Flask, request, render_template, send_file, redirect, url_for
from google.cloud import firestore
import os
import json
import base64
from docx import Document
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore

app = Flask(__name__)


# # Set the Firebase service account key
# os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "credentials.json"

# # Initialize Firestore client
# db = firestore.Client()

# Step 1: Load Firebase Credentials from Environment Variables
firebase_key_base64 = os.getenv("FIREBASE_KEY")  # Load from environment
if not firebase_key_base64:
    raise ValueError("FIREBASE_KEY environment variable is not set")

# Step 2: Decode the Base64 string and initialize Firebase
firebase_key = json.loads(base64.b64decode(firebase_key_base64).decode("utf-8"))
cred = credentials.Certificate(firebase_key)
firebase_admin.initialize_app(cred)

# Step 3: Initialize Firestore
db = firestore.client()

# Helper function to replace placeholders in the document
def replace_text_in_run(run, key, value):
    if key in run.text:
        run.text = run.text.replace(key, value)
        run.font.bold = True

def replace_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in placeholders.items():
                replace_text_in_run(run, key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in placeholders.items():
                            replace_text_in_run(run, key, value)

def get_template_path(oltype, internship):
    """Select the template path based on the offer letter type (oltype) or internship."""
    if internship:
        return "NEW_INTERN.docx"
    templates = {
        "BDA": "NEW_JOB-5LPA.docx",
        "Senior": "NEW_JOB-7LPA.docx",
        "Graphic Designer/Human Resource": "NEW_JOB-2.2LPA.docx",
        "Telecaller/Catalog": "NEW_JOB-1.8LPA.docx",
    }
    return templates.get(oltype, None)

def create_offer_letter(data, template_path, output_folder, unique_id):
    # Load the offer letter template
    doc = Document(template_path)
    
    # Prepare placeholders
    placeholders = {
        "S_NO": data.get("serial_number", "N/A"),  # Use Firestore unique ID
        "[NAME]": data.get("name", "N/A"),
        "FATHER": data.get("father_name", "N/A"),
        "MOBILE": data.get("mobile", "N/A"),
        "DATE": data.get("DATE", "N/A"),
        "JODA": data.get("joda", "N/A"),
        "CITY": data.get("city", "N/A"),
        "AADHAR": data.get("aadhar", "N/A"),
        "PAN": data.get("pan", "N/A"),
        "<ROLE>": data.get("role", "N/A"),
        "<MANAGER>": data.get("manager", "N/A"),
        "ITFP": data.get("type", "N/A"),  # Internship type (Full-time/Part-time)
        "STIPEND": data.get("stipend", "N/A"),  # Stipend
    }
    
    # Replace placeholders in the document
    replace_placeholders(doc, placeholders)
    
    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Save the offer letter
    file_label = "INTERNSHIP" if data.get("internship") == "on" else "JOB"
    file_name = f"{file_label}_HETERIZE_INFOTECH_{data.get('serial_number', 'N/A')}_{data.get('name', 'N/A')}.docx"
    file_path = os.path.join(output_folder, file_name)
    doc.save(file_path)
    
    return file_path

def save_to_firestore(data):
    """Save the data into Firestore."""
    collection_name = "OFFER LETTER"
    
    # Add a timestamp
    data["timestamp"] = datetime.utcnow().isoformat()
    
    # Save the data in Firestore with a unique document ID
    doc_ref = db.collection(collection_name).document()
    unique_id = doc_ref.id  # Get Firestore-generated unique ID
    data["unique_id"] = unique_id  # Add unique ID to the data
    
    doc_ref.set(data)
    return unique_id

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Extract form data
        data = request.form.to_dict()
        
        # Determine if internship checkbox is checked
        internship = data.get("internship") == "on"
        
        # Determine the template based on the selected oltype or internship
        oltype = data.get("oltype")
        template_path = get_template_path(oltype, internship)
        if not template_path:
            return "Invalid offer letter type selected.", 400
        
        # Set the output folder
        output_folder = "offer_letters"
        
        try:
            # Save the data to Firestore and get the unique ID
            unique_id = save_to_firestore(data)
            
            # Generate the offer letter
            file_path = create_offer_letter(data, template_path, output_folder, unique_id)
            
            # Redirect to download the file
            return redirect(url_for("download", file_name=os.path.basename(file_path)))
        except Exception as e:
            return f"An error occurred: {e}", 500

    return render_template("index.html")

@app.route("/download/<file_name>")
def download(file_name):
    file_path = os.path.join("offer_letters", file_name)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "File not found", 404

if __name__ == "__main__":
    app.run(debug=True)
